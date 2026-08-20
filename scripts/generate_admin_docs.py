#!/usr/bin/env python3
"""Generate ENGAGE Admin Panel documentation as a Word document."""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor, Cm
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
OUT_DOC = ROOT / "ENGAGE_Admin_Panel_Documentation.docx"
DIAG_DIR = ROOT / "scripts" / "_doc_diagrams"


# ---------------------------------------------------------------------------
# Diagram helpers
# ---------------------------------------------------------------------------

def _font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
        "/Library/Fonts/Arial.ttf",
    ]
    for path in candidates:
        if os.path.exists(path):
            try:
                return ImageFont.truetype(path, size)
            except Exception:
                continue
    return ImageFont.load_default()


def _rounded_rect(draw, xy, radius, fill, outline=None, width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def _center_text(draw, box, text, font, fill=(255, 255, 255)):
    x0, y0, x1, y1 = box
    bbox = draw.textbbox((0, 0), text, font=font)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.text(((x0 + x1 - tw) / 2, (y0 + y1 - th) / 2), text, font=font, fill=fill)


def _multiline_center(draw, box, lines, font, fill=(255, 255, 255), gap=4):
    x0, y0, x1, y1 = box
    heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        widths.append(bbox[2] - bbox[0])
        heights.append(bbox[3] - bbox[1])
    total_h = sum(heights) + gap * (len(lines) - 1)
    cy = (y0 + y1 - total_h) / 2
    for i, line in enumerate(lines):
        tw = widths[i]
        draw.text(((x0 + x1 - tw) / 2, cy), line, font=font, fill=fill)
        cy += heights[i] + gap


def _arrow(draw, start, end, color=(55, 65, 81), width=3):
    draw.line([start, end], fill=color, width=width)
    x0, y0 = start
    x1, y1 = end
    import math
    angle = math.atan2(y1 - y0, x1 - x0)
    size = 10
    p1 = (x1 - size * math.cos(angle - 0.4), y1 - size * math.sin(angle - 0.4))
    p2 = (x1 - size * math.cos(angle + 0.4), y1 - size * math.sin(angle + 0.4))
    draw.polygon([end, p1, p2], fill=color)


def create_architecture_diagram(path: Path):
    w, h = 1400, 900
    img = Image.new("RGB", (w, h), (248, 250, 252))
    draw = ImageDraw.Draw(img)
    title_f = _font(28, True)
    box_f = _font(16, True)
    small_f = _font(13)
    draw.text((40, 24), "ENGAGE Admin — High-Level Architecture", font=title_f, fill=(15, 23, 42))

    # Browser / Admin App
    _rounded_rect(draw, (40, 80, 560, 860), 16, (255, 255, 255), (37, 99, 235), 3)
    draw.text((60, 95), "ENGAGE Admin (React + Vite)", font=box_f, fill=(30, 64, 175))

    modules = [
        (60, 140, "Login"),
        (220, 140, "Sidebar Shell"),
        (60, 230, "Spaces"),
        (220, 230, "Meeting Rooms"),
        (380, 230, "Room Status"),
        (60, 320, "Bookings"),
        (220, 320, "Companies"),
        (380, 320, "Users"),
        (60, 410, "Wallet Report"),
        (220, 410, "Transactions"),
    ]
    for x, y, label in modules:
        _rounded_rect(draw, (x, y, x + 140, y + 55), 10, (37, 99, 235), None, 0)
        _center_text(draw, (x, y, x + 140, y + 55), label, small_f)

    # Context
    _rounded_rect(draw, (60, 500, 500, 580), 10, (239, 246, 255), (59, 130, 246), 2)
    _multiline_center(draw, (60, 500, 500, 580), ["UserContext + localStorage", "token • user • role (admin | manager)"], small_f, (30, 64, 175))

    # API layer label
    _rounded_rect(draw, (60, 620, 500, 700), 10, (236, 253, 245), (16, 185, 129), 2)
    _multiline_center(draw, (60, 620, 500, 700), ["Axios API Client", "Bearer token interceptor"], small_f, (6, 95, 70))

    # External systems
    _rounded_rect(draw, (680, 140, 1320, 420), 16, (255, 255, 255), (5, 150, 105), 3)
    draw.text((700, 155), "Backend API", font=box_f, fill=(6, 95, 70))
    _rounded_rect(draw, (720, 210, 1280, 310), 12, (16, 185, 129), None, 0)
    _multiline_center(draw, (720, 210, 1280, 310), ["engage-app.astererp.com/api/v1", "Auth • Locations • Rooms • Bookings", "Companies • Users • Wallets"], small_f)

    _rounded_rect(draw, (680, 480, 1320, 700), 16, (255, 255, 255), (245, 158, 11), 3)
    draw.text((700, 495), "Firebase Storage", font=box_f, fill=(146, 64, 14))
    _rounded_rect(draw, (720, 560, 1280, 650), 12, (245, 158, 11), None, 0)
    _multiline_center(draw, (720, 560, 1280, 650), ["Image uploads for Spaces", "and Meeting Rooms"], small_f)

    _arrow(draw, (520, 660), (680, 280))
    _arrow(draw, (520, 660), (680, 580))
    draw.text((540, 400), "HTTPS", font=small_f, fill=(55, 65, 81))
    draw.text((540, 620), "Upload", font=small_f, fill=(55, 65, 81))

    img.save(path)


def create_login_flow_diagram(path: Path):
    w, h = 1200, 780
    img = Image.new("RGB", (w, h), (248, 250, 252))
    draw = ImageDraw.Draw(img)
    title_f = _font(26, True)
    f = _font(14, True)
    s = _font(12)
    draw.text((40, 20), "Authentication Flow", font=title_f, fill=(15, 23, 42))

    steps = [
        (80, 90, "Admin / Manager\nenters email &\npassword"),
        (320, 90, "POST /auth/login\nvia authApi"),
        (560, 90, "API returns\ntoken + user"),
        (800, 90, "Store token &\nuser in\nlocalStorage"),
        (320, 320, "Role check:\nadmin or\nmanager?"),
        (80, 520, "Access Denied\n(alert)"),
        (560, 520, "setUserData()\nUserContext"),
        (800, 520, "Navigate to\n/home/locations"),
    ]
    colors = [
        (37, 99, 235), (5, 150, 105), (5, 150, 105), (37, 99, 235),
        (245, 158, 11), (220, 38, 38), (37, 99, 235), (16, 185, 129),
    ]
    boxes = []
    for i, (x, y, text) in enumerate(steps):
        box = (x, y, x + 180, y + 110)
        boxes.append(box)
        _rounded_rect(draw, box, 12, colors[i], None, 0)
        _multiline_center(draw, box, text.split("\n"), s)

    # Arrows
    _arrow(draw, (260, 145), (320, 145))
    _arrow(draw, (500, 145), (560, 145))
    _arrow(draw, (740, 145), (800, 145))
    _arrow(draw, (890, 200), (890, 280))
    draw.line([(890, 280), (410, 280)], fill=(55, 65, 81), width=3)
    _arrow(draw, (410, 280), (410, 320))
    _arrow(draw, (320, 430), (200, 520))  # no
    _arrow(draw, (500, 375), (560, 520))  # yes to setUser
    _arrow(draw, (740, 575), (800, 575))

    draw.text((250, 470), "No", font=f, fill=(220, 38, 38))
    draw.text((470, 450), "Yes", font=f, fill=(16, 185, 129))

    img.save(path)


def create_booking_flow_diagram(path: Path):
    w, h = 1300, 820
    img = Image.new("RGB", (w, h), (248, 250, 252))
    draw = ImageDraw.Draw(img)
    title_f = _font(26, True)
    s = _font(13)
    draw.text((40, 20), "Booking Creation Flow", font=title_f, fill=(15, 23, 42))

    # Two entry paths
    _rounded_rect(draw, (60, 90, 320, 180), 12, (37, 99, 235), None, 0)
    _multiline_center(draw, (60, 90, 320, 180), ["Bookings page", "Add Booking"], s)
    _rounded_rect(draw, (420, 90, 760, 180), 12, (37, 99, 235), None, 0)
    _multiline_center(draw, (420, 90, 760, 180), ["Meeting Room Status", "Click free slot / Make Booking"], s)

    _rounded_rect(draw, (240, 240, 580, 340), 12, (59, 130, 246), None, 0)
    _multiline_center(draw, (240, 240, 580, 340), ["Booking Form", "title • date • time • location", "room • company • user"], s)

    _arrow(draw, (190, 180), (340, 240))
    _arrow(draw, (590, 180), (480, 240))

    _rounded_rect(draw, (240, 400, 580, 490), 12, (5, 150, 105), None, 0)
    _multiline_center(draw, (240, 400, 580, 490), ["POST /bookings", "status = confirmed"], s)
    _arrow(draw, (410, 340), (410, 400))

    # Outcomes
    _rounded_rect(draw, (60, 560, 300, 680), 12, (16, 185, 129), None, 0)
    _multiline_center(draw, (60, 560, 300, 680), ["Success", "Credits deducted", "Slot reserved"], s)
    _rounded_rect(draw, (360, 560, 600, 680), 12, (220, 38, 38), None, 0)
    _multiline_center(draw, (360, 560, 600, 680), ["403 Insufficient", "wallet credits"], s)
    _rounded_rect(draw, (660, 560, 900, 680), 12, (245, 158, 11), None, 0)
    _multiline_center(draw, (660, 560, 900, 680), ["409 Conflict", "Room already booked"], s)

    _arrow(draw, (320, 490), (180, 560))
    _arrow(draw, (410, 490), (480, 560))
    _arrow(draw, (500, 490), (780, 560))

    # Cancel path
    _rounded_rect(draw, (980, 240, 1240, 340), 12, (100, 116, 139), None, 0)
    _multiline_center(draw, (980, 240, 1240, 340), ["Cancel Booking", "POST /bookings/cancel/:id"], s)
    _rounded_rect(draw, (980, 400, 1240, 490), 12, (16, 185, 129), None, 0)
    _multiline_center(draw, (980, 400, 1240, 490), ["Status → cancelled", "Credits may restore"], s)
    _arrow(draw, (1110, 340), (1110, 400))

    img.save(path)


def create_data_flow_diagram(path: Path):
    w, h = 1400, 900
    img = Image.new("RGB", (w, h), (248, 250, 252))
    draw = ImageDraw.Draw(img)
    title_f = _font(26, True)
    box_f = _font(15, True)
    s = _font(12)
    draw.text((40, 20), "System Data Flow", font=title_f, fill=(15, 23, 42))

    # Layers
    layers = [
        (40, 80, 1360, 200, (239, 246, 255), (37, 99, 235), "Presentation Layer", "React Pages & Modals • Sidebar • Dialogs • Forms"),
        (40, 240, 1360, 360, (236, 253, 245), (5, 150, 105), "Application Layer", "Feature modules • UserContext • Axios API modules"),
        (40, 400, 1360, 560, (255, 247, 237), (245, 158, 11), "Integration Layer", "REST API (engage-app.astererp.com)  •  Firebase Storage"),
        (40, 600, 1360, 860, (254, 242, 242), (220, 38, 38), "Domain Data", ""),
    ]
    for x0, y0, x1, y1, fill, border, title, subtitle in layers:
        _rounded_rect(draw, (x0, y0, x1, y1), 14, fill, border, 3)
        draw.text((x0 + 24, y0 + 16), title, font=box_f, fill=border)
        if subtitle:
            draw.text((x0 + 24, y0 + 50), subtitle, font=s, fill=(55, 65, 81))

    # Domain entities
    entities = [
        (80, 680, "Location"),
        (280, 680, "Meeting Room"),
        (520, 680, "Company"),
        (720, 680, "User"),
        (900, 680, "Wallet"),
        (1080, 680, "Booking"),
        (80, 780, "Amenity"),
        (280, 780, "Transaction"),
        (520, 780, "Report"),
    ]
    for x, y, label in entities:
        _rounded_rect(draw, (x, y, x + 170, y + 50), 10, (220, 38, 38), None, 0)
        _center_text(draw, (x, y, x + 170, y + 50), label, s)

    # Vertical arrows between layers
    for x in (200, 700, 1100):
        _arrow(draw, (x, 200), (x, 240), (100, 116, 139), 2)
        _arrow(draw, (x, 360), (x, 400), (100, 116, 139), 2)
        _arrow(draw, (x, 560), (x, 600), (100, 116, 139), 2)

    img.save(path)


def create_entity_relationship_diagram(path: Path):
    w, h = 1300, 780
    img = Image.new("RGB", (w, h), (248, 250, 252))
    draw = ImageDraw.Draw(img)
    title_f = _font(26, True)
    box_f = _font(14, True)
    s = _font(11)
    draw.text((40, 20), "Entity Relationship Overview", font=title_f, fill=(15, 23, 42))

    entities = {
        "Location": (80, 120, ["id", "name", "city", "address", "status", "image"]),
        "MeetingRoom": (480, 80, ["id", "name", "floor", "creditsPerSlot", "openingTime", "amenities[]"]),
        "Company": (80, 380, ["id", "name", "email", "LocationId", "gstn", "status"]),
        "User": (480, 380, ["id", "name", "role", "company_id", "email", "status"]),
        "Wallet": (860, 380, ["id", "meeting_room_credits", "printing_credits"]),
        "Booking": (860, 80, ["id", "title", "date", "start/end", "status", "user_id"]),
        "Transaction": (860, 580, ["id", "amount", "type", "reason", "createdAt"]),
    }

    boxes = {}
    for name, (x, y, fields) in entities.items():
        box_h = 40 + 22 * len(fields)
        box = (x, y, x + 240, y + box_h)
        boxes[name] = box
        _rounded_rect(draw, box, 10, (255, 255, 255), (37, 99, 235), 2)
        _rounded_rect(draw, (x, y, x + 240, y + 34), 10, (37, 99, 235), None, 0)
        # fix bottom corners of header by overlaying
        draw.rectangle((x, y + 20, x + 240, y + 34), fill=(37, 99, 235))
        _center_text(draw, (x, y, x + 240, y + 34), name, box_f)
        cy = y + 42
        for field in fields:
            draw.text((x + 14, cy), field, font=s, fill=(51, 65, 85))
            cy += 22

    def link(a, b, label="1:N"):
        ax0, ay0, ax1, ay1 = boxes[a]
        bx0, by0, bx1, by1 = boxes[b]
        start = ((ax0 + ax1) / 2, ay1)
        end = ((bx0 + bx1) / 2, by0)
        # smarter midpoints for side links
        if abs(ax0 - bx0) > 100 and abs(ay0 - by0) < 80:
            start = (ax1, (ay0 + ay1) / 2)
            end = (bx0, (by0 + by1) / 2)
        elif ay0 > by0 + 50:
            start = ((ax0 + ax1) / 2, ay0)
            end = ((bx0 + bx1) / 2, by1)
        _arrow(draw, start, end, (100, 116, 139), 2)
        mx, my = ((start[0] + end[0]) / 2, (start[1] + end[1]) / 2 - 12)
        draw.text((mx, my), label, font=s, fill=(71, 85, 105))

    # Manual arrows for clearer layout
    _arrow(draw, (320, 200), (480, 160), (100, 116, 139), 2)
    draw.text((360, 140), "1:N rooms", font=s, fill=(71, 85, 105))
    _arrow(draw, (200, 340), (200, 380), (100, 116, 139), 2)
    draw.text((210, 350), "1:N companies", font=s, fill=(71, 85, 105))
    _arrow(draw, (320, 450), (480, 430), (100, 116, 139), 2)
    draw.text((360, 420), "1:N users", font=s, fill=(71, 85, 105))
    _arrow(draw, (720, 430), (860, 430), (100, 116, 139), 2)
    draw.text((760, 400), "1:1", font=s, fill=(71, 85, 105))
    _arrow(draw, (720, 200), (860, 160), (100, 116, 139), 2)
    draw.text((760, 140), "1:N bookings", font=s, fill=(71, 85, 105))
    _arrow(draw, (600, 530), (920, 200), (100, 116, 139), 2)
    draw.text((700, 280), "user books", font=s, fill=(71, 85, 105))
    _arrow(draw, (980, 530), (980, 580), (100, 116, 139), 2)
    draw.text((990, 545), "1:N", font=s, fill=(71, 85, 105))

    img.save(path)


def create_module_nav_diagram(path: Path):
    w, h = 1100, 700
    img = Image.new("RGB", (w, h), (248, 250, 252))
    draw = ImageDraw.Draw(img)
    title_f = _font(26, True)
    s = _font(13)
    draw.text((40, 20), "Navigation & Module Map", font=title_f, fill=(15, 23, 42))

    _rounded_rect(draw, (420, 80, 680, 150), 12, (15, 23, 42), None, 0)
    _center_text(draw, (420, 80, 680, 150), "ENGAGE Sidebar", s)

    items = [
        ("Space", "/home/locations", (37, 99, 235)),
        ("Meeting Rooms", "/home/meeting-rooms", (37, 99, 235)),
        ("Meeting Room Status", "/home/meeting-room-status", (37, 99, 235)),
        ("Bookings", "/home/bookings", (37, 99, 235)),
        ("Users", "/home/users", (5, 150, 105)),
        ("Company", "/home/community", (5, 150, 105)),
        ("Report", "/home/reports", (245, 158, 11)),
        ("Logout", "/", (220, 38, 38)),
    ]
    y = 200
    for label, route, color in items:
        _rounded_rect(draw, (120, y, 420, y + 48), 10, color, None, 0)
        _center_text(draw, (120, y, 420, y + 48), label, s)
        _rounded_rect(draw, (480, y, 980, y + 48), 10, (255, 255, 255), color, 2)
        _center_text(draw, (480, y, 980, y + 48), route, s, fill=(30, 41, 59))
        _arrow(draw, (420, y + 24), (480, y + 24), color, 2)
        y += 60

    img.save(path)


def create_wallet_flow_diagram(path: Path):
    w, h = 1200, 700
    img = Image.new("RGB", (w, h), (248, 250, 252))
    draw = ImageDraw.Draw(img)
    title_f = _font(26, True)
    s = _font(13)
    draw.text((40, 20), "Wallet & Credits Data Flow", font=title_f, fill=(15, 23, 42))

    steps = [
        (80, 100, "Users page\nUpdate Wallet", (37, 99, 235)),
        (360, 100, "PUT /company/\nwallets/:id", (5, 150, 105)),
        (640, 100, "Credits updated\n(meeting room /\nprinting)", (16, 185, 129)),
        (80, 320, "Create Booking", (37, 99, 235)),
        (360, 320, "API deducts\ncreditsPerSlot\n× slots", (245, 158, 11)),
        (640, 320, "Wallet\nTransaction\n(debit)", (220, 38, 38)),
        (80, 520, "Report page\nuser + date range", (37, 99, 235)),
        (360, 520, "POST /company/\nwallet/report", (5, 150, 105)),
        (640, 520, "Export Excel\nor CSV", (100, 116, 139)),
        (920, 320, "View history\n/transactions/:id", (59, 130, 246)),
    ]
    for x, y, text, color in steps:
        _rounded_rect(draw, (x, y, x + 220, y + 110), 12, color, None, 0)
        _multiline_center(draw, (x, y, x + 220, y + 110), text.split("\n"), s)

    _arrow(draw, (300, 155), (360, 155))
    _arrow(draw, (580, 155), (640, 155))
    _arrow(draw, (300, 375), (360, 375))
    _arrow(draw, (580, 375), (640, 375))
    _arrow(draw, (860, 375), (920, 375))
    _arrow(draw, (300, 575), (360, 575))
    _arrow(draw, (580, 575), (640, 575))

    img.save(path)


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(15, 23, 42)
    return h


def add_para(doc, text, size=11, bold=False, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def set_cell_shading(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        set_cell_shading(hdr[i], "1E3A8A")
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = str(val)
            for p in cells[c_i].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9)
            if r_i % 2 == 1:
                set_cell_shading(cells[c_i], "F1F5F9")
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_image(doc, path: Path, width=6.5):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()


def add_page_number(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ENGAGE Admin Panel Documentation  |  Page ")
    set_run_font(run, size=9, color=(100, 116, 139))

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run2 = p.add_run()
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)
    set_run_font(run2, size=9, color=(100, 116, 139))


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

def build():
    DIAG_DIR.mkdir(parents=True, exist_ok=True)

    diagrams = {
        "architecture": DIAG_DIR / "architecture.png",
        "login": DIAG_DIR / "login_flow.png",
        "booking": DIAG_DIR / "booking_flow.png",
        "dataflow": DIAG_DIR / "data_flow.png",
        "er": DIAG_DIR / "entity_relationship.png",
        "nav": DIAG_DIR / "module_nav.png",
        "wallet": DIAG_DIR / "wallet_flow.png",
    }
    create_architecture_diagram(diagrams["architecture"])
    create_login_flow_diagram(diagrams["login"])
    create_booking_flow_diagram(diagrams["booking"])
    create_data_flow_diagram(diagrams["dataflow"])
    create_entity_relationship_diagram(diagrams["er"])
    create_module_nav_diagram(diagrams["nav"])
    create_wallet_flow_diagram(diagrams["wallet"])

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ----- Cover -----
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ENGAGE")
    set_run_font(run, size=36, bold=True, color=(30, 58, 138))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Admin Panel Documentation")
    set_run_font(run, size=26, bold=True, color=(15, 23, 42))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Features • Application Flows • Data Flow Diagrams\n"
        "Product: Aster ENGAGE  |  Package: daftarkhwan_vms_admin\n"
        "Version 1.0  |  July 2026"
    )
    set_run_font(run, size=12, color=(71, 85, 105))

    doc.add_page_break()

    # ----- TOC-like outline -----
    add_heading_styled(doc, "Table of Contents", 1)
    toc_items = [
        "1. Introduction & Overview",
        "2. Technology Stack",
        "3. System Architecture",
        "4. Navigation & Module Map",
        "5. Authentication & Authorization",
        "6. Feature Modules (Detailed)",
        "7. Application Flow Diagrams",
        "8. Data Flow Diagrams",
        "9. Data Models & Entity Relationships",
        "10. API Inventory",
        "11. Role-Based Access Matrix",
        "12. Integrations",
        "13. Known Limitations",
    ]
    for item in toc_items:
        add_para(doc, item, size=12, space_after=4)
    doc.add_page_break()

    # ----- 1. Introduction -----
    add_heading_styled(doc, "1. Introduction & Overview", 1)
    add_para(
        doc,
        "ENGAGE Admin is a web-based operations console for managing coworking spaces, "
        "meeting rooms, tenant companies, users, bookings, and wallet credits. Although "
        "the repository package name is daftarkhwan_vms_admin, the product is not a classic "
        "visitor check-in VMS — it is a space and meeting-room administration panel for Aster ENGAGE.",
    )
    add_para(doc, "Primary goals of the admin panel:", bold=True)
    for item in [
        "Configure and maintain physical locations (spaces) and meeting rooms.",
        "Manage tenant companies and their users.",
        "Create, monitor, and cancel meeting room bookings.",
        "View real-time room slot availability for a selected date.",
        "Maintain wallet credits (meeting room / printing) and export transaction reports.",
    ]:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font(run, size=11)

    add_para(doc, "Backend API base URL: https://engage-app.astererp.com/api/v1/", bold=True)
    add_para(doc, "Default post-login landing page: /home/locations (Space module).")

    # ----- 2. Tech stack -----
    add_heading_styled(doc, "2. Technology Stack", 1)
    add_table(
        doc,
        ["Layer", "Technology"],
        [
            ["UI Framework", "React 19"],
            ["Build Tool", "Vite 7"],
            ["Routing", "react-router v7 (BrowserRouter)"],
            ["Styling", "Tailwind CSS v4 + feature CSS"],
            ["HTTP Client", "Axios (Bearer token interceptor)"],
            ["Auth Persistence", "localStorage (token, user)"],
            ["State Management", "React Context (UserContext) — no Redux"],
            ["Image Storage", "Firebase Storage"],
            ["Exports", "xlsx + file-saver (Excel / CSV)"],
            ["Select Controls", "react-select"],
        ],
        [2.2, 4.5],
    )

    # ----- 3. Architecture -----
    add_heading_styled(doc, "3. System Architecture", 1)
    add_para(
        doc,
        "The admin panel is a single-page application (SPA). Feature pages live under a shared "
        "Home layout that renders a persistent sidebar and an Outlet for nested routes. API "
        "communication is centralized through Axios modules under src/api/. Image uploads for "
        "locations and rooms go to Firebase Storage; all other business data flows through the ENGAGE REST API.",
    )
    add_image(doc, diagrams["architecture"], 6.4)

    add_heading_styled(doc, "3.1 Project Structure", 2)
    add_para(
        doc,
        "src/\n"
        "  api/           — Axios client and domain API modules\n"
        "  components/    — Shared UI (sidebar, dialogs, popups, loader)\n"
        "  context/       — UserContext (auth user / role)\n"
        "  features/      — Feature modules (auth, locations, rooms, bookings, users, wallet, …)\n"
        "  firebase/      — Firebase app initialization (Storage)\n"
        "  router/        — Route definitions\n"
        "  App.jsx        — BrowserRouter + UserProvider + routes",
        size=10,
    )

    # ----- 4. Navigation -----
    add_heading_styled(doc, "4. Navigation & Module Map", 1)
    add_para(
        doc,
        "The left sidebar branded “ENGAGE” provides primary navigation. Transaction history "
        "opens in a separate route (/transactions/:id) without the sidebar shell.",
    )
    add_image(doc, diagrams["nav"], 5.8)
    add_table(
        doc,
        ["Menu Label", "Route", "Feature Module"],
        [
            ["Space", "/home/locations", "Locations management"],
            ["Meeting Rooms", "/home/meeting-rooms", "Room configuration"],
            ["Meeting Room Status", "/home/meeting-room-status", "Slot grid & booking"],
            ["Bookings", "/home/bookings", "Booking list & CRUD"],
            ["Users", "/home/users", "User & wallet management"],
            ["Company", "/home/community", "Tenant companies"],
            ["Report", "/home/reports", "Wallet transaction report"],
            ["Logout", "/", "Clears storage and redirects"],
            ["(Transactions)", "/transactions/:id", "Per-wallet history (new tab)"],
        ],
        [2.0, 2.4, 2.3],
    )

    # ----- 5. Auth -----
    add_heading_styled(doc, "5. Authentication & Authorization", 1)
    add_heading_styled(doc, "5.1 Login Flow", 2)
    add_para(
        doc,
        "Users authenticate with email and password. Only roles admin and manager are allowed "
        "into the console. On success, the API authorization token and user object are stored "
        "in localStorage. Every subsequent Axios request attaches Authorization: Bearer <token>.",
    )
    add_image(doc, diagrams["login"], 6.2)

    add_heading_styled(doc, "5.2 Logout", 2)
    add_para(
        doc,
        "Logout clears localStorage, sessionStorage, and cookies for the current domain, then "
        "redirects to the login page (/).",
    )

    add_heading_styled(doc, "5.3 Route Protection Note", 2)
    add_para(
        doc,
        "There is no dedicated client-side ProtectedRoute guard on /home/* routes. Access is "
        "enforced at login (role gate) and via per-action UI checks for admin-only mutations. "
        "API calls without a valid token fail at the backend.",
    )

    # ----- 6. Features -----
    add_heading_styled(doc, "6. Feature Modules (Detailed)", 1)

    add_heading_styled(doc, "6.1 Space (Locations)", 2)
    add_para(
        doc,
        "Manage physical coworking sites. Admins can create, edit, and delete locations; "
        "managers can view and search. Images are uploaded to Firebase Storage.",
    )
    add_para(doc, "Capabilities:", bold=True)
    for item in [
        "List locations with client-side search (name, city, contact, email) and pagination (10/page).",
        "Add / Edit / Delete location (admin only).",
        "Fields: name, contactNumber, email, legalBusinessName, address, city, lat, lng, status, image.",
    ]:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font(run, size=11)

    add_heading_styled(doc, "6.2 Meeting Rooms", 2)
    add_para(
        doc,
        "Configure bookable rooms under locations, including pricing in credits, capacity, "
        "operating hours, available weekdays, and amenities.",
    )
    for item in [
        "List rooms with search and filters (location, floor).",
        "Add / Edit / Delete room (admin only); Firebase image upload.",
        "Fields: name, locationId, floor, creditsPerSlot, pricePerCredit, seatingCapacity, "
        "opening/closing time, status, amenities (multi-select), availableDays, image.",
        "Amenities loaded from GET /amenities.",
    ]:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font(run, size=11)

    add_heading_styled(doc, "6.3 Meeting Room Status", 2)
    add_para(
        doc,
        "Operational day view showing 30-minute slots for a selected room and date. Free slots "
        "can be booked; booked slots show details and allow cancellation. Past slots for the "
        "current day are non-bookable. Hours follow the room’s opening/closing times (default 09:00–21:30).",
    )
    for item in [
        "Select location → room → date → Check Status.",
        "Click free slot or “Make a Booking” to open booking form (preset fields).",
        "Click booked slot to view title, description, company, user, time — then cancel if needed.",
    ]:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font(run, size=11)

    add_heading_styled(doc, "6.4 Bookings", 2)
    add_para(
        doc,
        "Paginated booking management with server-side paging (page_size=10). Supports create "
        "and cancel. Dependent dropdowns load rooms by location and users by company.",
    )
    for item in [
        "List, search (client filter on current page), location filter, sort by date.",
        "Create booking payload: title, description, date, startTime/endTime, location_id, "
        "room_id, company_id, user_id, status=confirmed.",
        "Cancel via POST /bookings/cancel/:id.",
        "Handled errors: 403 insufficient wallet credits; 409 room conflict.",
    ]:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font(run, size=11)

    add_heading_styled(doc, "6.5 Company", 2)
    add_para(
        doc,
        "Manage tenant companies associated with locations. Admins can add and edit companies. "
        "A delete icon exists in the list UI but the parent page does not currently wire onDelete.",
    )
    for item in [
        "List, search, filter by location/status.",
        "Fields: name, email, contactNumber, businessType, websiteUrl, LocationId, locationName, "
        "reference, billingEmail, gstn, status.",
    ]:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font(run, size=11)

    add_heading_styled(doc, "6.6 Users", 2)
    add_para(
        doc,
        "Manage platform users and their credit wallets. Roles assignable at create: admin, "
        "manager, member.",
    )
    for item in [
        "List and search users.",
        "Add / Edit user (admin only).",
        "Update Wallet — adjust meeting room credits and printing credits.",
        "Open transaction history in a new tab (/transactions/:walletId).",
        "Delete confirmation currently removes from local UI state only (no delete API call).",
        "Fields: name, role, email, company_id, phoneNumber, password (create), credit_types "
        "(prepaid/postpaid), auto_renew, status.",
    ]:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font(run, size=11)

    add_heading_styled(doc, "6.7 Report (Wallet Transactions)", 2)
    add_para(
        doc,
        "Generate a filtered wallet transaction report by selecting a user and date range, then "
        "export results to Excel or CSV.",
    )
    for item in [
        "POST /company/wallet/report with user and date filters.",
        "Table display of transactions (amount, type, reason, timestamps, metadata).",
        "Export via xlsx / file-saver.",
    ]:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font(run, size=11)

    # ----- 7. App flows -----
    add_heading_styled(doc, "7. Application Flow Diagrams", 1)

    add_heading_styled(doc, "7.1 Booking Creation & Cancellation", 2)
    add_para(
        doc,
        "Bookings can be started from the Bookings page or from Meeting Room Status. The form "
        "submits to POST /bookings. The backend validates wallet balance and room availability.",
    )
    add_image(doc, diagrams["booking"], 6.3)

    add_heading_styled(doc, "7.2 Wallet & Credits Flow", 2)
    add_para(
        doc,
        "Credits are topped up from the Users page, consumed when bookings are created, and "
        "audited via transaction history and the Report module.",
    )
    add_image(doc, diagrams["wallet"], 6.2)

    add_heading_styled(doc, "7.3 Typical Operator Journey", 2)
    add_para(doc, "End-to-end path for day-to-day operations:", bold=True)
    for i, step in enumerate(
        [
            "Sign in as admin or manager.",
            "Ensure Space (location) and Meeting Rooms are configured (admin).",
            "Ensure Company and Users exist with sufficient wallet credits.",
            "Open Meeting Room Status — pick location, room, date — review free/booked slots.",
            "Create a booking from a free slot or from the Bookings page.",
            "If needed, cancel a booking from the list or status detail modal.",
            "Review wallet activity under Report or per-user Transactions.",
        ],
        start=1,
    ):
        p = doc.add_paragraph(f"{i}. {step}")
        for run in p.runs:
            set_run_font(run, size=11)

    # ----- 8. Data flow -----
    add_heading_styled(doc, "8. Data Flow Diagrams", 1)
    add_para(
        doc,
        "Data moves from UI presentation through feature/application modules into the "
        "integration layer (REST API and Firebase), and persists as domain entities on the backend.",
    )
    add_image(doc, diagrams["dataflow"], 6.4)

    add_heading_styled(doc, "8.1 Request Lifecycle", 2)
    for item in [
        "User action in a React page/modal updates local component state.",
        "Handler calls a domain API function (e.g., addNewBooking).",
        "Axios interceptor attaches Bearer token from localStorage.",
        "Request sent to https://engage-app.astererp.com/api/v1/<resource>.",
        "Response JSON is mapped into component state (lists, modals, toasts).",
        "For images: file uploaded to Firebase Storage; returned URL included in create/update payload.",
    ]:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            set_run_font(run, size=11)

    add_heading_styled(doc, "8.2 Cross-Entity Data Dependencies", 2)
    add_table(
        doc,
        ["UI Action", "Dependent Loads", "Write Target"],
        [
            ["Create room", "Locations, Amenities", "POST /meeting-rooms (+ Firebase image)"],
            ["Create company", "Locations", "POST /company"],
            ["Create user", "Companies", "POST /auth/register/"],
            ["Create booking", "Locations → Rooms; Companies → Users", "POST /bookings"],
            ["Room status check", "Locations → Rooms; Bookings by room+date", "GET bookings/by-room-and-date"],
            ["Update wallet", "User wallet id", "PUT /company/wallets/:id"],
            ["Wallet report", "Users list", "POST /company/wallet/report"],
        ],
        [1.8, 2.5, 2.4],
    )

    # ----- 9. Models -----
    add_heading_styled(doc, "9. Data Models & Entity Relationships", 1)
    add_image(doc, diagrams["er"], 6.3)

    add_heading_styled(doc, "9.1 Entity Field Reference", 2)

    add_para(doc, "Location", bold=True)
    add_para(doc, "id, name, contactNumber, email, legalBusinessName, address, city, lat, lng, status, image/imageUrl")

    add_para(doc, "Meeting Room", bold=True)
    add_para(
        doc,
        "id, name, LocationId/location_id, location (nested), floor, creditsPerSlot, pricePerCredit, "
        "seatingCapacity, openingTime, closingTime, status, availableDays[], amenities[], image, availableSlotsCount",
    )

    add_para(doc, "Company", bold=True)
    add_para(
        doc,
        "id, name, email, contactNumber, businessType, websiteUrl, LocationId, locationName, location, "
        "reference, billingEmail, gstn, status",
    )

    add_para(doc, "User", bold=True)
    add_para(
        doc,
        "id, name, email, role (admin|manager|member), company_id, Company {name, locationName}, "
        "phoneNumber, password, credit_types, auto_renew, status, Wallet {id, meeting_room_credits, printing_credits}",
    )

    add_para(doc, "Booking", bold=True)
    add_para(
        doc,
        "id, title, description, date, startTime, endTime, location_id, room_id, company_id, user_id, "
        "status (confirmed|cancelled), nested Room, User, User.Company",
    )

    add_para(doc, "Wallet Transaction", bold=True)
    add_para(
        doc,
        "id, amount, type (credit|debit), reason, createdAt, metadata (JSON: company, date, startTime, "
        "roomName, location, slots, totalCredits), performedByUser, user",
    )

    # ----- 10. API -----
    add_heading_styled(doc, "10. API Inventory", 1)
    add_para(doc, "All endpoints are relative to https://engage-app.astererp.com/api/v1/", size=10)

    add_table(
        doc,
        ["Module", "Method", "Endpoint", "Purpose"],
        [
            ["Auth", "POST", "/auth/login", "Sign in"],
            ["Auth", "GET", "/auth/users/company/:id", "Users by company"],
            ["Users", "GET", "/auth/users", "List users"],
            ["Users", "POST", "/auth/register/", "Create user"],
            ["Users", "PUT", "/auth/users/edit/:id", "Edit user"],
            ["Locations", "GET", "/locations", "List locations"],
            ["Locations", "POST", "/locations", "Create location"],
            ["Locations", "PUT", "/locations/:id", "Update location"],
            ["Locations", "DELETE", "/locations/delete/", "Delete location"],
            ["Rooms", "GET", "/meeting-rooms", "List rooms"],
            ["Rooms", "GET", "/meeting-rooms/location/:id", "Rooms by location"],
            ["Rooms", "POST", "/meeting-rooms", "Create room"],
            ["Rooms", "PUT", "/meeting-rooms/:id", "Update room"],
            ["Rooms", "DELETE", "meeting-rooms/delete/:id", "Delete room"],
            ["Rooms", "GET", "/amenities", "List amenities"],
            ["Bookings", "GET", "/bookings?page&page_size", "Paginated bookings"],
            ["Bookings", "POST", "/bookings", "Create booking"],
            ["Bookings", "POST", "/bookings/cancel/:id", "Cancel booking"],
            ["Bookings", "GET", "/bookings/by-room-and-date", "Slots for room/date"],
            ["Bookings", "GET", "/bookings/:id", "Booking detail"],
            ["Company", "GET", "/company", "List companies"],
            ["Company", "POST", "/company", "Create company"],
            ["Company", "PUT", "/company/:id", "Update company"],
            ["Company", "GET", "/company/location-id/:id", "Companies by location"],
            ["Wallet", "PUT", "/company/wallets/:id", "Update wallet credits"],
            ["Wallet", "GET", "/company/wallet/transactions/:id", "Transaction history"],
            ["Wallet", "POST", "/company/wallet/report", "Filtered wallet report"],
        ],
        [1.2, 0.9, 2.6, 2.0],
    )

    # ----- 11. RBAC -----
    add_heading_styled(doc, "11. Role-Based Access Matrix", 1)
    add_table(
        doc,
        ["Capability", "Admin", "Manager", "Member"],
        [
            ["Login to admin panel", "Yes", "Yes", "No"],
            ["View spaces / rooms / bookings / companies / users / reports", "Yes", "Yes", "—"],
            ["Add / Edit / Delete locations", "Yes", "No (UI blocked)", "—"],
            ["Add / Edit / Delete meeting rooms", "Yes", "No (UI blocked)", "—"],
            ["Add / Edit companies", "Yes", "No (UI blocked)", "—"],
            ["Add / Edit / Delete users", "Yes", "No (UI blocked)", "—"],
            ["Create / Cancel bookings", "Yes", "Yes", "—"],
            ["Meeting room status book/cancel", "Yes", "Yes", "—"],
            ["Update wallet / run reports", "Yes", "Yes", "—"],
        ],
        [3.2, 1.2, 1.6, 1.0],
    )
    add_para(
        doc,
        "Note: Role checks for mutations are primarily enforced in the UI. Backend authorization "
        "may apply additional rules.",
        size=10,
    )

    # ----- 12. Integrations -----
    add_heading_styled(doc, "12. Integrations", 1)
    add_table(
        doc,
        ["Integration", "Usage"],
        [
            ["Aster ENGAGE REST API", "Primary backend for all business operations"],
            ["Firebase Storage", "Image uploads for locations and meeting rooms (path: images/<filename>)"],
            ["xlsx + file-saver", "Export wallet reports to Excel/CSV"],
            ["react-select", "Multi-select amenities on meeting rooms"],
        ],
        [2.4, 4.3],
    )
    add_para(
        doc,
        "Payments, email, and SMS providers are not integrated in this frontend. Classic visitor "
        "management (check-in/check-out) is not present in this codebase.",
    )

    # ----- 13. Limitations -----
    add_heading_styled(doc, "13. Known Limitations", 1)
    for item in [
        "No client-side route guard on /home/* — deep links load the shell without re-checking auth.",
        "User delete confirms in UI but does not call a delete API.",
        "Company delete control is present in the list but not wired from the parent page.",
        "src/features/community/community.jsx is an unused legacy companies UI; the router uses companies.jsx.",
        "Manager vs admin restrictions are UI-gated; document backend enforcement separately if needed.",
        "Package naming (daftarkhwan_vms_admin) reflects historical naming; product branding is ENGAGE.",
    ]:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font(run, size=11)

    add_heading_styled(doc, "Document Control", 1)
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Document title", "ENGAGE Admin Panel Documentation"],
            ["Audience", "Product, engineering, operations"],
            ["Source codebase", "daftarkhwan_vms_admin"],
            ["Generated", "July 2026"],
            ["Format", "Microsoft Word (.docx)"],
        ],
        [2.0, 4.7],
    )

    add_page_number(doc)
    doc.save(str(OUT_DOC))
    print(f"Wrote {OUT_DOC}")
    return OUT_DOC


if __name__ == "__main__":
    build()
